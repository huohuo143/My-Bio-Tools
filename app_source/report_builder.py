"""Generate Word, Excel and ZIP deliverables for rice gene analysis."""

from __future__ import annotations

from collections import defaultdict
from datetime import datetime, timezone
import csv
import io
import json
from pathlib import Path
import re
import zipfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

from rice_gene_core import (
    AnalysisBundle,
    SEQUENCE_TYPES,
    prediction_consistency,
    safe_file_stem,
    sequence_records_to_fasta,
)
from rice_efp import EFP_DATA_SOURCES, EFP_SOURCE_GLOSSARY, EFP_GUIDE_URL, EFP_URL, efp_source_display_label, expression_top_rows
from report_interpretation import build_rule_interpretations


# Use the actual macOS/PostScript family name so Word and LibreOffice resolve
# the requested FangSong face instead of rendering missing-glyph boxes.
CHINESE_FONT = "华文仿宋"
WESTERN_FONT = "Times New Roman"
EXCEL_FONT = "Arial"
BODY_SIZE = 10.5
TITLE_SIZE = 18
HEADING1_SIZE = 16
HEADING2_SIZE = 13
ACCENT = RGBColor(15, 118, 110)
MUTED = RGBColor(102, 112, 133)
TABLE_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
CAUTION_FILL = "FFF7E6"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def _set_rfonts(element: OxmlElement, size: float | None = None, bold: bool | None = None, latin_font: str = WESTERN_FONT) -> None:
    r_pr = element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for theme_attribute in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        r_fonts.attrib.pop(qn(theme_attribute), None)
    r_fonts.set(qn("w:ascii"), latin_font)
    r_fonts.set(qn("w:hAnsi"), latin_font)
    r_fonts.set(qn("w:eastAsia"), CHINESE_FONT)
    r_fonts.set(qn("w:cs"), latin_font)
    if size is not None:
        half_points = str(int(round(size * 2)))
        for tag in ("w:sz", "w:szCs"):
            node = r_pr.find(qn(tag))
            if node is None:
                node = OxmlElement(tag)
                r_pr.append(node)
            node.set(qn("w:val"), half_points)
    if bold is not None:
        b = r_pr.find(qn("w:b"))
        if b is None:
            b = OxmlElement("w:b")
            r_pr.append(b)
        b.set(qn("w:val"), "1" if bold else "0")


def _format_run(run, size: float = BODY_SIZE, bold: bool | None = None, color: RGBColor | None = None) -> None:
    # Keep mixed Chinese/Latin runs script-aware: Word uses eastAsia for CJK
    # glyphs and ascii/hAnsi for Latin letters, numbers, IDs and sequences.
    run.font.name = WESTERN_FONT
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    _set_rfonts(run._element, size=size, bold=bold, latin_font=WESTERN_FONT)


def _configure_styles(doc: Document) -> None:
    style_tokens = {
        "Normal": (BODY_SIZE, False, RGBColor(23, 32, 51), 0, 6),
        "Title": (TITLE_SIZE, True, ACCENT, 0, 12),
        "Heading 1": (HEADING1_SIZE, True, ACCENT, 18, 10),
        "Heading 2": (HEADING2_SIZE, True, ACCENT, 14, 7),
        "Heading 3": (12, True, RGBColor(23, 32, 51), 10, 5),
    }
    for name, (size, bold, color, before, after) in style_tokens.items():
        style = doc.styles[name]
        style.font.name = WESTERN_FONT
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = color
        _set_rfonts(style.element, size=size, bold=bold)
        style_p_pr = style.element.get_or_add_pPr()
        style_borders = style_p_pr.find(qn("w:pBdr"))
        if style_borders is not None:
            style_p_pr.remove(style_borders)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        if name.startswith("Heading"):
            style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = WESTERN_FONT
        style.font.size = Pt(BODY_SIZE)
        _set_rfonts(style.element, size=BODY_SIZE, bold=False)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def _remove_paragraph_borders(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is not None:
        p_pr.remove(borders)


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def _set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _format_paragraph_runs(
    paragraph,
    size: float = BODY_SIZE,
    bold: bool | None = None,
    color: RGBColor | None = None,
) -> None:
    for run in paragraph.runs:
        _format_run(run, size=size, bold=bold, color=color)


def _add_key_value_table(doc: Document, rows: list[tuple[str, object]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    _set_table_geometry(table, [2700, 6660])
    table.rows[0].cells[0].text = "项目"
    table.rows[0].cells[1].text = "内容"
    for cell in table.rows[0].cells:
        cell._tc.get_or_add_tcPr().append(_cell_shading(TABLE_FILL))
    _set_repeat_table_header(table.rows[0])
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = "" if value is None else str(value)
        _set_table_geometry(table, [2700, 6660])
    for row_index, row in enumerate(table.rows):
        _prevent_row_split(row)
        for col_index, cell in enumerate(row.cells):
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _format_paragraph_runs(paragraph, bold=row_index == 0 or col_index == 0)


def _cell_shading(fill: str) -> OxmlElement:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    return shd


def _add_rows_table(doc: Document, rows: list[dict[str, object]], columns: list[tuple[str, str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    _set_table_geometry(table, widths)
    for cell, (_, label) in zip(table.rows[0].cells, columns):
        cell.text = label
        cell._tc.get_or_add_tcPr().append(_cell_shading(TABLE_FILL))
    _set_repeat_table_header(table.rows[0])
    for item in rows:
        cells = table.add_row().cells
        for cell, (key, _) in zip(cells, columns):
            value = item.get(key, "")
            cell.text = "" if value is None else str(value)
        _set_table_geometry(table, widths)
    for row_index, row in enumerate(table.rows):
        _prevent_row_split(row)
        for cell in row.cells:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _format_paragraph_runs(paragraph, size=9.5, bold=row_index == 0)


def _add_callout(doc: Document, label: str, text: str, *, fill: str = CALLOUT_FILL) -> None:
    """Add one concise evidence/interpretation box using explicit Word geometry."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    _set_table_geometry(table, [TABLE_WIDTH_DXA])
    _set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    cell._tc.get_or_add_tcPr().append(_cell_shading(fill))
    _set_cell_margins(cell, top=130, start=180, bottom=130, end=180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    lead = paragraph.add_run(f"{label}  ")
    _format_run(lead, size=10, bold=True, color=ACCENT)
    body = paragraph.add_run(text)
    _format_run(body, size=10, color=RGBColor(23, 32, 51))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def _interpretation_rows(bundle: AnalysisBundle, sections: set[str] | None = None) -> list[dict[str, object]]:
    rows = bundle.interpretations or build_rule_interpretations(bundle)
    return [row for row in rows if sections is None or str(row.get("section") or "") in sections]


def _add_interpretation_blocks(doc: Document, bundle: AnalysisBundle, sections: set[str]) -> None:
    for row in _interpretation_rows(bundle, sections):
        ai_assisted = str(row.get("section") or "").startswith("ai_")
        label = str(row.get("title") or "结果解读")
        if ai_assisted:
            label += "（AI辅助推断·待人工核验）"
        text = (
            f"{row.get('interpretation') or ''}\n"
            f"证据依据：{row.get('evidence_basis') or '—'}；"
            f"证据等级：{row.get('evidence_level') or '—'}；"
            f"置信度：{row.get('confidence') or '—'}。\n"
            f"解读边界：{row.get('limitations') or '—'}\n"
            f"建议下一步：{row.get('recommended_action') or '—'}"
        )
        _add_callout(doc, label, text, fill=CAUTION_FILL if ai_assisted else CALLOUT_FILL)


def _compact_text(value: object, limit: int = 90) -> str:
    text = " ".join(str(value or "").split())
    return text if len(text) <= limit else f"{text[: limit - 1].rstrip()}…"


def _format_fraction(value: object) -> str:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return "—"
    return f"{number:.1%}" if 0 <= number <= 1 else f"{number:.3g}"


def _feature_bp(rows: list[dict[str, object]], feature_types: set[str]) -> int:
    intervals: list[tuple[int, int]] = []
    for row in rows:
        if str(row.get("feature_type") or "") not in feature_types:
            continue
        try:
            start, end = sorted((int(row["start"]), int(row["end"])))
        except (KeyError, TypeError, ValueError):
            continue
        intervals.append((start, end))
    if not intervals:
        return 0
    intervals.sort()
    merged = [intervals[0]]
    for start, end in intervals[1:]:
        previous_start, previous_end = merged[-1]
        if start <= previous_end + 1:
            merged[-1] = (previous_start, max(previous_end, end))
        else:
            merged.append((start, end))
    return sum(end - start + 1 for start, end in merged)


def _transcript_summary_rows(bundle: AnalysisBundle) -> list[dict[str, object]]:
    features: dict[tuple[str, str], list[dict[str, object]]] = defaultdict(list)
    for row in bundle.gene_features:
        features[(str(row.get("rap_gene") or ""), str(row.get("transcript_id") or ""))].append(row)
    summaries: list[dict[str, object]] = []
    for row in bundle.transcript_models:
        if row.get("status") != "matched" or not row.get("transcript_id"):
            continue
        key = (str(row.get("rap_gene") or ""), str(row.get("transcript_id") or ""))
        tx_features = features.get(key, [])
        exon_count = len({str(item.get("feature_id") or item.get("feature_number") or "") for item in tx_features if item.get("feature_type") == "exon"})
        five_bp = _feature_bp(tx_features, {"5UTR", "five_prime_UTR"})
        three_bp = _feature_bp(tx_features, {"3UTR", "three_prime_UTR"})
        strand = "+" if int(row.get("strand") or 0) == 1 else "−" if int(row.get("strand") or 0) == -1 else "?"
        summaries.append(
            {
                "transcript_id": row.get("transcript_id"),
                "canonical": "Yes" if row.get("is_canonical") else "No",
                "strand_label": strand,
                "exon_count": exon_count or "—",
                "cds_bp": _feature_bp(tx_features, {"CDS"}) or "—",
                "utr_bp": f"5′ {five_bp or '—'} / 3′ {three_bp or '—'}",
            }
        )
    return summaries


def _variant_priority_rows(rows: list[dict[str, object]], limit: int = 15) -> list[dict[str, object]]:
    keywords = ("stop_gained", "frameshift", "splice", "missense", "inframe", "synonymous")

    def priority(row: dict[str, object]) -> tuple[int, int]:
        consequence = str(row.get("coding_consequence") or "").casefold()
        score = next((index for index, keyword in enumerate(keywords) if keyword in consequence), len(keywords))
        try:
            position = int(row.get("position") or 0)
        except (TypeError, ValueError):
            position = 0
        return score, position

    output = []
    for row in sorted(rows, key=priority)[:limit]:
        output.append(
            {
                "position": row.get("position") or "—",
                "alleles": f"{row.get('ref') or '—'} → {row.get('alt') or '—'}",
                "region": _compact_text(row.get("region") or "unclassified", 32),
                "consequence": _compact_text(row.get("coding_consequence") or "not annotated", 72),
                "af": _format_fraction(row.get("allele_frequency")),
                "qc": _compact_text(row.get("ref_validation") or row.get("status") or "—", 22),
            }
        )
    return output


def _add_chart_figures(
    doc: Document,
    charts: dict[str, bytes],
    prefix: str,
    caption_prefix: str,
) -> int:
    png_names = sorted(name for name in charts if name.startswith(prefix) and name.endswith(".png"))
    for index, name in enumerate(png_names, start=1):
        caption = doc.add_paragraph()
        caption.paragraph_format.keep_with_next = True
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run(f"Figure {caption_prefix}.{index}  {_friendly_chart_title(name)}")
        _format_run(run, size=9, bold=True, color=RGBColor(51, 65, 85))
        doc.add_picture(io.BytesIO(charts[name]), width=Inches(6.25))
        picture_paragraph = doc.paragraphs[-1]
        _set_picture_alt(picture_paragraph, f"{caption_prefix}.{index} {name.rsplit('/', 1)[-1]}")
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_paragraph.paragraph_format.keep_together = True
        picture_paragraph.paragraph_format.space_after = Pt(6)
    return len(png_names)


def _friendly_chart_title(name: str) -> str:
    """Translate internal artifact filenames into reader-facing figure captions."""
    stem = name.rsplit("/", 1)[-1].removesuffix(".png")
    exact = {
        "variant_overview": "Variant landscape and regional composition",
        "haplotype_frequency": "Haplotype frequency distribution",
        "haplotype_group_heatmap": "Haplotype frequencies across population groups",
    }
    if stem in exact:
        return exact[stem]
    if stem.startswith("protein_domains_"):
        protein = stem.removeprefix("protein_domains_").replace("_", " ")
        return f"Integrated protein architecture · {protein}"
    if stem.startswith("gene_structure_"):
        gene = stem.removeprefix("gene_structure_").replace("_", " ")
        return f"Transcript-oriented gene models · {gene}"
    if stem.startswith("tfbs_"):
        transcript = stem.removeprefix("tfbs_").replace("_", " ")
        return f"Promoter TFBS landscape · {transcript}"
    if stem.startswith("combined_"):
        protein = stem.removeprefix("combined_").replace("_", " ")
        return f"Integrated protein-localization prediction · {protein}"
    if stem.startswith("bar_"):
        payload = stem.removeprefix("bar_")
        for source, label in (("_ricestress_rma", "Stress atlas (RMA)"), ("_rice_rma", "Developmental atlas (RMA)"), ("_ricestress_mas", "Stress atlas (MAS)"), ("_rice_mas", "Developmental atlas (MAS)")):
            if payload.endswith(source):
                return f"Rice eFP expression profile · {payload.removesuffix(source)} · {label}"
        return f"Rice eFP expression profile · {payload.replace('_', ' ')}"
    if stem.startswith("heatmap_"):
        return f"Rice eFP expression matrix · {stem.removeprefix('heatmap_').replace('_', ' ')}"
    return stem.replace("_", " ")


def _set_picture_alt(paragraph, text: str) -> None:
    """Set Word-native title/description on every drawing in a paragraph."""
    for doc_properties in paragraph._p.iter(qn("wp:docPr")):
        doc_properties.set("title", text)
        doc_properties.set("descr", text)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    _format_run(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    tail = paragraph.add_run(" 页")
    _format_run(tail, size=9, color=MUTED)


def _start_chapter(doc: Document, number: str, title: str, purpose: str) -> None:
    doc.add_page_break()
    heading = doc.add_heading(f"{number}. {title}", level=1)
    heading.paragraph_format.keep_with_next = True
    _add_callout(doc, "本章目的", purpose)


def _add_table_or_status(
    doc: Document,
    rows: list[dict[str, object]],
    columns: list[tuple[str, str]],
    widths: list[int],
    empty_text: str,
) -> bool:
    if not rows:
        _add_callout(doc, "状态", empty_text, fill=CAUTION_FILL)
        return False
    _add_rows_table(doc, rows, columns, widths)
    return True


def _add_first_png(
    doc: Document,
    charts: dict[str, bytes],
    prefixes: tuple[str, ...],
    figure_number: str,
    caption: str,
) -> bool:
    names = [name for name in sorted(charts) if name.endswith(".png") and any(name.startswith(prefix) for prefix in prefixes)]
    if not names:
        return False
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"Figure {figure_number}  {caption}")
    _format_run(run, size=9, bold=True, color=RGBColor(51, 65, 85))
    doc.add_picture(io.BytesIO(charts[names[0]]), width=Inches(6.25))
    picture = doc.paragraphs[-1]
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.paragraph_format.keep_together = True
    picture.paragraph_format.space_after = Pt(6)
    _set_picture_alt(picture, caption)
    return True


def _evidence_reference_rows(bundle: AnalysisBundle) -> list[dict[str, object]]:
    reference_by_doi = {
        str(row.get("doi") or "").casefold(): row
        for row in bundle.ricedata_references
        if row.get("doi")
    }
    rows: list[dict[str, object]] = []
    directly_used: set[str] = set()
    for evidence in bundle.genetic_evidence:
        linked = [value.strip() for value in str(evidence.get("linked_dois") or "").split(",") if value.strip()]
        directly_used.update(value.casefold() for value in linked)
        titles = [str(reference_by_doi.get(value.casefold(), {}).get("title") or value) for value in linked]
        rows.append({
            "evidence": _compact_text(evidence.get("evidence_text"), 135),
            "paper": _compact_text("; ".join(titles) or "未解析到关联论文", 100),
            "doi": ", ".join(linked) or "—",
            "source": evidence.get("source_type") or "—",
            "verification": evidence.get("verification_status") or "—",
        })
    for reference in bundle.ricedata_references:
        doi = str(reference.get("doi") or "")
        if doi.casefold() in directly_used:
            continue
        rows.append({
            "evidence": "RiceData 关联文献（未直接映射到当前遗传证据）",
            "paper": _compact_text(reference.get("title"), 100),
            "doi": doi or "—",
            "source": f"RiceData ref {reference.get('reference_id') or '—'}",
            "verification": reference.get("verification_status") or "需全文核验",
        })
    return rows


def build_word_report_legacy(
    bundle: AnalysisBundle,
    include_full_sequences: bool,
    efp_charts: dict[str, bytes] | None = None,
    prediction_charts: dict[str, bytes] | None = None,
    deep_charts: dict[str, bytes] | None = None,
) -> bytes:
    doc = Document()
    _configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.text = "My Bio Tools · 水稻基因一站式分析"
    _format_paragraph_runs(header, size=9)
    header.runs[0].font.color.rgb = MUTED
    _add_page_number(section.footer.paragraphs[0])

    title = doc.add_paragraph(style="Title")
    _remove_paragraph_borders(title)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("水稻基因一站式分析报告")
    _format_paragraph_runs(title, size=TITLE_SIZE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("RAP/MSU · IRGSP-1.0 · 可追溯计算预测")
    _format_run(run, size=10.5, color=MUTED)
    subtitle.paragraph_format.space_after = Pt(14)

    doc.add_heading("1. 分析概览", level=1)
    _add_key_value_table(
        doc,
        [
            ("分析模式", bundle.mode),
            ("输入类型", bundle.input_type),
            ("输入数量", len(bundle.inputs)),
            ("序列记录", len(bundle.sequences)),
            ("预测记录", len(bundle.predictions)),
            ("RiceData 记录", len(bundle.ricedata_rows)),
            ("eFP 表达记录", len(bundle.efp_rows)),
            ("生成时间", bundle.generated_at),
        ],
    )
    _add_callout(
        doc,
        "阅读提示",
        "正文优先呈现图形、关键结论与高优先级记录；完整逐行数据保留在配套 Excel 和 ZIP 中，避免明细表淹没科研判断。",
    )

    doc.add_heading("2. ID 解析与映射", level=1)
    mapping_rows = bundle.mapping_rows or [{"input_id": item, "status": "not_resolved"} for item in bundle.inputs]
    _add_rows_table(
        doc,
        mapping_rows,
        [
            ("input_id", "输入"),
            ("input_type", "类型"),
            ("resolved_rap_gene", "RAP gene"),
            ("resolved_msu_id", "MSU ID"),
            ("status", "状态"),
        ],
        [1700, 1300, 2100, 2400, 1860],
    )

    doc.add_heading("3. RiceData 基因信息", level=1)
    _add_rows_table(
        doc,
        bundle.ricedata_rows or [{"status": "未选择或无可用记录"}],
        [
            ("check", "查询 ID"),
            ("GeneName", "Gene name"),
            ("GeneSymbol", "Symbol"),
            ("RAP_Locus", "RAP locus"),
            ("MSU_Locus", "MSU locus"),
            ("status", "状态"),
        ],
        [1500, 1900, 1100, 1800, 1960, 1100],
    )
    detailed_rows = [
        row
        for row in bundle.ricedata_rows
        if any(row.get(key) for key in ("突变体表型", "定位与克隆", "时空表达谱", "亚细胞定位", "生物学功能"))
    ]
    for row in detailed_rows[:20]:
        doc.add_heading(str(row.get("RAP_Locus") or row.get("MSU_Locus") or row.get("check") or "RiceData"), level=2)
        for key in ("突变体表型", "定位与克隆", "时空表达谱", "亚细胞定位", "生物学功能"):
            if row.get(key):
                paragraph = doc.add_paragraph()
                paragraph.add_run(f"{key}：").bold = True
                paragraph.add_run(str(row[key]))
                _format_paragraph_runs(paragraph)

    doc.add_heading("4. Rice eFP 表达谱", level=1)
    top_rows = expression_top_rows(bundle.efp_rows, limit=3)
    _add_rows_table(
        doc,
        top_rows or [{"status": "未选择或无可用记录"}],
        [
            ("msu_locus", "MSU locus"),
            ("data_source_label", "数据源"),
            ("rank", "排名"),
            ("tissue", "组织/处理"),
            ("expression_level", "Expression"),
            ("standard_deviation", "SD"),
        ],
        [1700, 1800, 700, 2600, 1300, 1260],
    )
    charts = dict(efp_charts or {})
    png_names = sorted(name for name in charts if name.endswith(".png") and name.startswith("heatmap_"))
    if not png_names:
        png_names = sorted(name for name in charts if name.endswith(".png") and name.startswith("bar_"))
    for index, name in enumerate(png_names, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.keep_with_next = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(f"Figure 4.{index}  {_friendly_chart_title(name)}")
        _format_paragraph_runs(paragraph, size=9, bold=True, color=RGBColor(51, 65, 85))
        doc.add_picture(io.BytesIO(charts[name]), width=Inches(6.35))
        picture_paragraph = doc.paragraphs[-1]
        _set_picture_alt(picture_paragraph, _friendly_chart_title(name))
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_paragraph.paragraph_format.keep_together = True
        picture_paragraph.paragraph_format.space_after = Pt(6)
    source_note = doc.add_paragraph()
    source_note.add_run(
        f"数据来源：{EFP_URL}。保留官网 Absolute 模式原始数值，不跨数据源重新标准化。"
    )
    _format_paragraph_runs(source_note)

    doc.add_heading("5. 序列结果", level=1)
    _add_rows_table(
        doc,
        [record.summary_row() for record in bundle.sequences] or [{"status": "无可用序列"}],
        [
            ("transcript_id", "Transcript/ID"),
            ("sequence_type", "序列类型"),
            ("length", "长度"),
            ("source", "来源"),
            ("status", "状态"),
        ],
        [2500, 1700, 900, 2660, 1600],
    )

    doc.add_heading("6. 蛋白定位预测", level=1)
    _add_rows_table(
        doc,
        [result.summary_row() for result in bundle.predictions] or [{"status": "未选择预测"}],
        [
            ("protein_id", "蛋白"),
            ("tool", "工具"),
            ("classification", "分类"),
            ("provider", "服务来源"),
            ("status", "状态"),
        ],
        [1900, 1700, 1900, 1900, 1960],
    )
    prediction_figure_index = 0
    for name, payload in sorted((prediction_charts or {}).items()):
        if not (name.startswith("combined_") and name.endswith(".png")):
            continue
        prediction_figure_index += 1
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.keep_with_next = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(f"Figure 6.{prediction_figure_index}  {_friendly_chart_title(name)}")
        _format_paragraph_runs(paragraph, size=9, bold=True, color=RGBColor(51, 65, 85))
        doc.add_picture(io.BytesIO(payload), width=Inches(6.35))
        picture_paragraph = doc.paragraphs[-1]
        _set_picture_alt(picture_paragraph, _friendly_chart_title(name))
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_paragraph.paragraph_format.keep_together = True
        picture_paragraph.paragraph_format.space_after = Pt(6)
    note = doc.add_paragraph()
    run = note.add_run("说明：以上结果均为 computational prediction，不代表已经获得实验定位证据。")
    _format_run(run, size=BODY_SIZE, bold=True, color=RGBColor(122, 90, 0))
    doc.add_heading("预测一致性说明", level=2)
    for item in prediction_consistency(bundle.predictions):
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)
        _format_paragraph_runs(paragraph)

    visual_charts = dict(deep_charts or {})

    doc.add_heading("7. 蛋白结构预测与功能位点", level=1)
    matched_domains = [row for row in bundle.protein_domains if row.get("status") == "matched"]
    domain_proteins = {str(row.get("protein_id")) for row in matched_domains if row.get("protein_id")}
    _add_callout(
        doc,
        "结构概览",
        f"覆盖 {len(domain_proteins)} 个蛋白模型；获得 {len(matched_domains)} 个结构域/家族/区域命中和 {len(bundle.functional_sites)} 个功能位点。图形用于定位关系，表格仅保留前 15 个代表性命中。",
    )
    if not _add_chart_figures(doc, visual_charts, "protein_domains/", "7"):
        _add_callout(doc, "图形状态", "本次没有可嵌入的蛋白结构图；请核对结构域服务结果与绘图后端警告。", fill=CAUTION_FILL)
    domain_rows = []
    for row in matched_domains[:15]:
        domain_rows.append(
            {
                "protein_id": row.get("protein_id"),
                "feature": _compact_text(row.get("name") or row.get("accession") or row.get("feature_type"), 45),
                "interval": f"{row.get('start') or '—'}–{row.get('end') or '—'} aa",
                "database": row.get("database") or "—",
                "accession": row.get("accession") or "—",
            }
        )
    _add_rows_table(
        doc,
        domain_rows or [{"feature": "未选择、无命中或外部服务未返回数据"}],
        [("protein_id", "Protein"), ("feature", "Feature"), ("interval", "Interval"), ("database", "Database"), ("accession", "Accession")],
        [1900, 2800, 1500, 1500, 1660],
    )
    if bundle.functional_sites:
        doc.add_heading("代表性功能位点", level=2)
        _add_rows_table(
            doc,
            bundle.functional_sites[:10],
            [("protein_id", "Protein"), ("description", "Site"), ("residue", "Residue"), ("start", "Position"), ("database", "Database")],
            [1900, 3400, 1000, 1200, 1860],
        )

    doc.add_heading("8. 基因结构与转录本", level=1)
    transcript_rows = _transcript_summary_rows(bundle)
    canonical_count = sum(row.get("canonical") == "Yes" for row in transcript_rows)
    _add_callout(
        doc,
        "模型概览",
        f"解析到 {len(transcript_rows)} 个 matched transcript，其中 canonical {canonical_count} 个。结构图统一按转录方向显示为 5′→3′，便于比较 exon、CDS 与 UTR。",
    )
    if not _add_chart_figures(doc, visual_charts, "gene_structure/", "8"):
        _add_callout(doc, "图形状态", "本次没有可嵌入的 exon/CDS/UTR 轨迹图；请核对 gene feature 坐标与绘图后端警告。", fill=CAUTION_FILL)
    _add_rows_table(
        doc,
        transcript_rows or [{"transcript_id": "未选择、无命中或外部服务未返回数据"}],
        [("transcript_id", "Transcript"), ("canonical", "Canonical"), ("strand_label", "Strand"), ("exon_count", "Exons"), ("cds_bp", "CDS bp"), ("utr_bp", "UTR bp")],
        [2500, 1100, 800, 900, 1300, 2760],
    )

    doc.add_heading("9. 启动子与上游调控", level=1)
    tf_families = {str(row.get("tf_family")) for row in bundle.promoter_tfbs if row.get("tf_family")}
    _add_callout(
        doc,
        "调控概览",
        f"共获得 {len(bundle.promoter_tfbs)} 个 motif-based TFBS，涉及 {len(tf_families)} 个 TF family；候选 TF 共 {len(bundle.upstream_tfs)} 个。黄色区域表示近端启动子（−500 至 TSS）。这些关系是计算预测，不代表已验证调控。",
        fill=CAUTION_FILL,
    )
    if not _add_chart_figures(doc, visual_charts, "promoter_regulation/", "9"):
        _add_callout(doc, "图形状态", "本次没有可嵌入的启动子 TFBS 图；请核对 PlantRegMap 结果与绘图后端警告。", fill=CAUTION_FILL)
    candidate_rows = []
    for row in bundle.upstream_tfs[:12]:
        candidate_rows.append(
            {
                "rank": row.get("rank") or "—",
                "tf": row.get("tf") or "—",
                "tf_family": row.get("tf_family") or "—",
                "hit_count": row.get("hit_count") or "—",
                "best_p_value": f"{float(row.get('best_p_value')):.2g}" if row.get("best_p_value") not in (None, "") else "—",
                "nearest_tss_bp": row.get("nearest_tss_bp") or "—",
            }
        )
    _add_rows_table(
        doc,
        candidate_rows or [{"tf": "未选择、无命中或外部服务未返回候选 TF"}],
        [("rank", "Rank"), ("tf", "TF/motif"), ("tf_family", "Family"), ("hit_count", "Hits"), ("best_p_value", "Best P"), ("nearest_tss_bp", "Nearest TSS (bp)")],
        [700, 2000, 1500, 900, 1500, 2760],
    )

    doc.add_heading("10. 自然变异与单倍型", level=1)
    _add_callout(
        doc,
        "变异概览",
        f"当前记录 {len(bundle.variants)} 个变异；形成 {len(bundle.haplotypes)} 个可计算单倍型。正文展示分布图、单倍型频率和前 15 个优先变异，完整变异表位于 Excel/ZIP。",
    )
    if not _add_chart_figures(doc, visual_charts, "variation/", "10"):
        _add_callout(doc, "图形状态", "本次没有可嵌入的变异/单倍型图；请核对 VCF 样本列、变异坐标与绘图后端警告。", fill=CAUTION_FILL)
    if bundle.haplotypes:
        haplotype_rows = []
        for row in bundle.haplotypes[:12]:
            haplotype_rows.append(
                {
                    "haplotype": row.get("haplotype") or "—",
                    "sample_count": row.get("sample_count") or 0,
                    "sample_frequency": _format_fraction(row.get("sample_frequency")),
                    "filtered_variant_count": row.get("filtered_variant_count") or "—",
                    "subgroup_frequency": _compact_text(row.get("subgroup_frequency") or "—", 70),
                }
            )
        doc.add_heading("单倍型频率摘要", level=2)
        _add_rows_table(
            doc,
            haplotype_rows,
            [("haplotype", "Haplotype"), ("sample_count", "Samples"), ("sample_frequency", "Frequency"), ("filtered_variant_count", "Sites"), ("subgroup_frequency", "Population distribution")],
            [1500, 1100, 1400, 900, 4460],
        )
    else:
        _add_callout(
            doc,
            "单倍型状态",
            "本次没有可计算的样本基因型矩阵，不能据数据库变异列表推断单倍型。请上传含 GT 样本列的 VCF；如需群体分层，再附 sample/group 对照表。",
            fill=CAUTION_FILL,
        )
    doc.add_heading("高优先级变异（Top 15）", level=2)
    _add_rows_table(
        doc,
        _variant_priority_rows(bundle.variants) or [{"position": "未选择、无命中或外部服务未返回数据"}],
        [("position", "Position"), ("alleles", "REF → ALT"), ("region", "Region"), ("consequence", "Consequence"), ("af", "ALT AF"), ("qc", "REF/QC")],
        [1100, 1300, 1800, 2860, 1000, 1300],
    )

    doc.add_heading("11. miRNA/RNAi 分析", level=1)
    _add_callout(
        doc,
        "证据边界",
        "以下结果均为计算预测；只有导入实验或经核验文献证据后，才能标记为已有证据。",
        fill=CAUTION_FILL,
    )
    _add_rows_table(
        doc,
        bundle.mirna_targets[:20] or [{"small_rna": "未选择、无命中或外部服务未返回数据"}],
        [("small_rna", "miRNA/sRNA"), ("target_transcript", "Target"), ("target_start", "Start"), ("expectation", "Expectation"), ("upe", "UPE"), ("evidence_status", "Evidence")],
        [1560] * 6,
    )
    _add_chart_figures(doc, visual_charts, "mirna_rnai/", "11")

    doc.add_heading("12. 文献与已知遗传证据", level=1)
    literature_rows = []
    for row in bundle.literature_rows[:20]:
        item = dict(row)
        item["title"] = _compact_text(row.get("title"), 88)
        item["evidence_tags"] = _compact_text(row.get("evidence_tags"), 42)
        literature_rows.append(item)
    _add_rows_table(
        doc,
        literature_rows or [{"title": "未选择、无命中或外部服务未返回数据"}],
        [("pmid", "PMID"), ("doi", "DOI"), ("year", "Year"), ("title", "Title"), ("evidence_tags", "Tags"), ("verification_status", "Verification")],
        [1000, 1500, 700, 3400, 1460, 1300],
    )
    _add_chart_figures(doc, visual_charts, "literature_evidence/", "12")

    doc.add_heading("13. 科研判断卡", level=1)
    cards = [
        ("数据库已知证据", f"RiceData/人工导入遗传证据 {len(bundle.genetic_evidence)} 条。"),
        ("计算支持", f"结构域 {len(bundle.protein_domains)}，TFBS {len(bundle.promoter_tfbs)}，miRNA 靶点 {len(bundle.mirna_targets)} 条。"),
        ("证据冲突/缺口", "请核对外部服务警告、REF/assembly 冲突以及文献全文。"),
        ("建议验证方向", "优先围绕高置信结构域/位点、近 TSS motif、功能变异和有文献线索的互作进行实验。"),
    ]
    _add_key_value_table(doc, cards)

    doc.add_heading("14. 来源、警告与可追溯信息", level=1)
    for warning in bundle.warnings or ["未记录额外警告。"]:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(warning)
        _format_paragraph_runs(paragraph)
    doc.add_heading("数据与服务来源", level=2)
    for source in bundle.sources:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(source)
        _format_paragraph_runs(paragraph)

    if include_full_sequences:
        doc.add_page_break()
        doc.add_heading("附录：完整序列", level=1)
        for sequence_type in SEQUENCE_TYPES:
            fasta = sequence_records_to_fasta(bundle.sequences, sequence_type)
            if not fasta:
                continue
            doc.add_heading(sequence_type, level=2)
            for line in fasta.rstrip().splitlines():
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                run = paragraph.add_run(line)
                _format_run(run, size=8.5)

    for paragraph in doc.paragraphs:
        if not paragraph.runs:
            paragraph.add_run("")
        if paragraph.style.name == "Heading 1":
            _format_paragraph_runs(paragraph, size=HEADING1_SIZE, bold=True)
        elif paragraph.style.name == "Heading 2":
            _format_paragraph_runs(paragraph, size=HEADING2_SIZE, bold=True)
        elif paragraph.style.name == "Title":
            _format_paragraph_runs(paragraph, size=TITLE_SIZE, bold=True)
        else:
            for run in paragraph.runs:
                if run.font.size is None:
                    _format_run(run)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def build_word_report(
    bundle: AnalysisBundle,
    include_full_sequences: bool,
    efp_charts: dict[str, bytes] | None = None,
    prediction_charts: dict[str, bytes] | None = None,
    deep_charts: dict[str, bytes] | None = None,
) -> bytes:
    """Build the v1.9.1 evidence-led report with seven coherent chapters."""
    doc = Document()
    _configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.text = "My Bio Tools v1.9.1 · 水稻基因一站式分析"
    _format_paragraph_runs(header, size=9, color=MUTED)
    _add_page_number(section.footer.paragraphs[0])

    title = doc.add_paragraph(style="Title")
    _remove_paragraph_borders(title)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("水稻基因一站式分析报告")
    _format_paragraph_runs(title, size=TITLE_SIZE, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _format_run(subtitle.add_run("基因身份 → 已有证据 → 表达 → 分子结构 → 调控与变异 → 综合判断"), size=10.5, color=MUTED)
    subtitle.paragraph_format.space_after = Pt(14)
    _add_key_value_table(doc, [
        ("分析对象", ", ".join(bundle.inputs)),
        ("分析模式", bundle.mode),
        ("参考边界", "IRGSP-1.0；RAP/MSU 结果按来源分别保留"),
        ("生成时间", bundle.generated_at),
        ("正文策略", "仅展示主图和 Top 记录；完整明细见 Excel/ZIP/附录"),
        ("结果解读", "离线科研规则" + (" + 大模型增强（待人工核验）" if bundle.interpretation_status.get("effective_mode") == "llm" else "")),
    ])
    doc.add_heading("报告解读摘要", level=2)
    _add_interpretation_blocks(doc, bundle, {"overall", "ai_overall"})

    _start_chapter(doc, "1", "基因概览与 ID 映射", "先确认输入基因的统一身份、RAP/MSU/Transcript 对应关系、GeneSymbol、assembly 和数据完整度。")
    matched_sequences = sum(record.status == "matched" for record in bundle.sequences)
    serious_warnings = [warning for warning in bundle.warnings if any(token in warning for token in ("assembly", "REF", "失败", "一对多", "不一致"))]
    _add_callout(doc, "关键结论", f"本次解析 {len(bundle.inputs)} 个输入，获得 {len(bundle.genetic_evidence)} 条遗传/功能证据、{len(bundle.literature_rows)} 篇文献、{len(bundle.efp_rows)} 条表达记录和 {matched_sequences} 条有效序列；需要优先注意 {len(serious_warnings)} 项可能改变解释的警告。")
    identity_rows = []
    for mapping in bundle.mapping_rows or [{"input_id": value, "status": "not_resolved"} for value in bundle.inputs]:
        rice = next((row for row in bundle.ricedata_rows if str(row.get("RAP_Locus") or "").casefold() == str(mapping.get("resolved_rap_gene") or "").casefold()), {})
        identity_rows.append({
            "input_id": mapping.get("input_id"),
            "rap": mapping.get("resolved_rap_gene") or rice.get("RAP_Locus") or "—",
            "msu": mapping.get("resolved_msu_id") or rice.get("MSU_Locus") or "—",
            "symbol": rice.get("GeneSymbol") or "—",
            "name_status": f"{rice.get('GeneName') or '名称未返回'} · {mapping.get('status') or '—'}",
        })
    _add_rows_table(doc, identity_rows, [("input_id", "Input / transcript"), ("rap", "RAP gene"), ("msu", "MSU locus"), ("symbol", "Symbol"), ("name_status", "Gene name / status")], [1600, 1900, 1900, 1500, 2460])
    _add_callout(doc, "注意事项", "RAP 与 MSU 属于不同注释体系。下游结果保留各自来源和 assembly；一对多映射不会静默合并。", fill=CAUTION_FILL)

    _start_chapter(doc, "2", "已知功能、突变体与关联文献", "把数据库功能描述、突变体/定位证据及其支持论文放在同一处，直接展示证据—文献对应关系。")
    linked_count = sum(bool(row.get("linked_dois")) for row in bundle.genetic_evidence)
    _add_callout(doc, "关键结论", f"共整理 {len(bundle.genetic_evidence)} 条遗传或功能证据，其中 {linked_count} 条已映射到 DOI；RiceData 另关联 {len(bundle.ricedata_references)} 篇论文。直接支持与仅关联文献分开标记。")
    evidence_rows = _evidence_reference_rows(bundle)
    _add_table_or_status(
        doc,
        evidence_rows[:20],
        [("evidence", "证据描述"), ("paper", "支持/关联论文"), ("doi", "DOI"), ("source", "来源"), ("verification", "核验状态")],
        [2500, 2100, 1500, 1200, 2060],
        "本次未取得可映射的遗传证据或关联文献；完整检索状态见第7章。",
    )
    _add_callout(doc, "注意事项", "“直接支持”表示数据库证据文本与引用年份/作者或 DOI 可对应；“RiceData 关联文献”只说明数据库页面建立了链接，具体机制仍需阅读全文核验。", fill=CAUTION_FILL)

    _start_chapter(doc, "3", "表达模式与生物学场景", "集中判断高表达组织、处理响应和本次选择的数据集边界，不在正文重复完整 eFP 明细。")
    top_rows = expression_top_rows(bundle.efp_rows, limit=3)
    selected_sources = list(dict.fromkeys(record.data_source for record in bundle.efp_rows))
    _add_callout(doc, "关键结论", f"本次覆盖 {len(selected_sources)} 个 eFP 数据源；正文展示每个基因/数据源的 Top 3 组织或处理。Absolute 表示同一数据集内的原始尺度丰度，不是 fold change，且不同数据源之间不能直接比较。")
    expression_charts = dict(efp_charts or {})
    if not _add_first_png(doc, expression_charts, ("heatmap_", "bar_"), "3.1", "Selected Rice eFP expression profile"):
        _add_callout(doc, "主图状态", "本次没有可嵌入的 eFP 主图；可能未选择表达分析、ID 未映射或外部服务未返回定量表。", fill=CAUTION_FILL)
    _add_table_or_status(doc, top_rows[:18], [("msu_locus", "MSU locus"), ("data_source_label", "数据源"), ("rank", "Rank"), ("tissue", "组织/处理"), ("expression_level", "Value"), ("standard_deviation", "SD")], [1800, 1800, 900, 2600, 1260, 1000], "本次没有可用的 eFP 定量记录。")
    _add_callout(doc, "注意事项", "RMA、MAS5 intensity 与官网未标明单位的 Expression Level 不能跨数据源比较。部分汇总型数据源的 SD 字段为 0，不代表没有细胞间或生物学变异；12 个数据源的来源、重复结构与完整边界见附录和 Excel。", fill=CAUTION_FILL)

    doc.add_heading("实验室已分析多组学", level=2)
    _add_callout(
        doc,
        "组学覆盖",
        f"命中 {len(bundle.lab_omics_datasets)} 个数据集、{len(bundle.lab_omics_differential)} 条差异记录和 {len(bundle.lab_omics_profiles)} 条项目内定量记录。主键为去model后缀的MSU locus；MSU model、RAP gene/model和原始ID逐条保留。",
    )
    lab_charts = dict(deep_charts or {})
    if not _add_first_png(doc, lab_charts, ("lab_omics/heatmap_cross_project_log2fc",), "3.2", "Wu Lab analysed multi-omics treatment-response heatmap"):
        _add_callout(doc, "多组学主图状态", "当前基因在首版实验室多组学库中没有合格记录，或登录授权尚未解锁数据库。", fill=CAUTION_FILL)
    lab_rows = []
    for row in bundle.lab_omics_differential[:24]:
        lab_rows.append({
            "msu_locus": row.get("msu_locus"),
            "dataset": _compact_text(row.get("dataset_name"), 32),
            "assay": row.get("assay"),
            "comparison": _compact_text(row.get("comparison_name"), 28),
            "log2fc": f"{float(row['log2fc']):.3f}" if row.get("log2fc") is not None else "—",
            "note": "描述性" if row.get("descriptive") else f"n={row.get('n_treatment') or '—'}+{row.get('n_control') or '—'}",
        })
    _add_table_or_status(
        doc,
        lab_rows,
        [("msu_locus", "MSU locus"), ("dataset", "Dataset"), ("assay", "Assay"), ("comparison", "Comparison"), ("log2fc", "log2FC"), ("note", "Repeat")],
        [1700, 2100, 1200, 2100, 900, 1360],
        "本次没有实验室多组学差异记录。",
    )
    _add_callout(doc, "多组学注意事项", "跨项目只使用各源分析表已有log2FC；项目内热图使用已有FPKM、TPM、count或归一化蛋白/PTM定量。不同组学的原始值不直接比较；缺失值为灰色；无可核实重复的数据只作描述性结果。", fill=CAUTION_FILL)
    doc.add_heading("多组学科研解读", level=2)
    _add_interpretation_blocks(doc, bundle, {"lab_omics", "ai_lab_omics"})

    _start_chapter(doc, "4", "序列、转录本与蛋白结构", "把输入 ID、RAP/MSU 映射、promoter/genomic/UTR/CDS/protein、真实基因结构和蛋白结构域合并到一条可追溯链。")
    cds_rows = [row for row in bundle.sequence_plot_rows if row.get("sequence_type") == "CDS"]
    consistent = sum(row.get("translation_consistency") == "consistent" for row in cds_rows)
    _add_callout(doc, "关键结论", f"共获得 {len(bundle.sequences)} 条序列记录；{consistent}/{len(cds_rows)} 条 CDS 在已选蛋白记录中通过精确翻译一致性检查。RAP/MSU 基因组长度或边界不同时分别展示，不做坐标强行叠加。")
    visual_charts = dict(deep_charts or {})
    if not _add_first_png(doc, visual_charts, ("sequence_structure/sequence_relationship_", "sequence_structure/sequence_availability_"), "4.1", "Input-to-sequence relationship and source-specific lengths"):
        _add_callout(doc, "主图状态", "本次没有可嵌入的序列关系图；序列明细仍保留在 Excel/ZIP。", fill=CAUTION_FILL)
    sequence_rows = []
    for record in bundle.sequences:
        translation_status = next(
            (
                row.get("translation_consistency")
                for row in bundle.sequence_plot_rows
                if row.get("transcript_id") == record.transcript_id
                and row.get("sequence_type") == record.sequence_type
                and row.get("source") == record.source
            ),
            "—",
        )
        if translation_status in {None, "", "not_applicable"}:
            translation_status = "—"
        sequence_rows.append(dict(record.summary_row(), translation_consistency=translation_status))
    _add_table_or_status(doc, sequence_rows[:16], [("transcript_id", "Transcript/ID"), ("sequence_type", "Type"), ("length", "Length"), ("source", "Source"), ("assembly", "Assembly"), ("translation_consistency", "Translation")], [2100, 1400, 900, 2500, 1200, 1260], "本次没有可用序列。")
    transcript_rows = _transcript_summary_rows(bundle)
    if transcript_rows:
        doc.add_heading("真实 exon/CDS/UTR 结构", level=2)
        _add_first_png(doc, visual_charts, ("gene_structure/",), "4.2", "Coordinate-aware exon, CDS and UTR models")
        _add_rows_table(doc, transcript_rows[:12], [("transcript_id", "Transcript"), ("canonical", "Canonical"), ("strand_label", "Strand"), ("exon_count", "Exons"), ("cds_bp", "CDS bp"), ("utr_bp", "UTR bp")], [2500, 1100, 800, 900, 1300, 2760])
    matched_domains = [row for row in bundle.protein_domains if row.get("status") == "matched"]
    if matched_domains:
        doc.add_heading("蛋白结构域与功能位点", level=2)
        _add_first_png(doc, visual_charts, ("protein_domains/",), "4.3", "Protein domain architecture")
        domain_rows = [{"protein_id": row.get("protein_id"), "feature": _compact_text(row.get("name") or row.get("accession") or row.get("feature_type"), 45), "interval": f"{row.get('start') or '—'}–{row.get('end') or '—'} aa", "database": row.get("database") or "—", "accession": row.get("accession") or "—"} for row in matched_domains[:12]]
        _add_rows_table(doc, domain_rows, [("protein_id", "Protein"), ("feature", "Feature"), ("interval", "Interval"), ("database", "Database"), ("accession", "Accession")], [1900, 2800, 1500, 1500, 1660])
    _add_callout(doc, "注意事项", "序列关系图中的条形长度仅在核苷酸或蛋白类别内归一化；真实 exon/CDS/UTR 坐标以基因结构图和 Gene_Features sheet 为准。", fill=CAUTION_FILL)

    _start_chapter(doc, "5", "定位、调控与遗传变异", "依次回答蛋白可能作用位置、上游调控线索以及序列变异可能带来的影响。")
    _add_callout(doc, "关键结论", f"定位预测 {len(bundle.predictions)} 条，启动子 TFBS {len(bundle.promoter_tfbs)} 条，miRNA/sRNA 靶点 {len(bundle.mirna_targets)} 条，变异 {len(bundle.variants)} 个，单倍型 {len(bundle.haplotypes)} 个。除数据库明确证据外，本章均为计算支持。")
    localization_rows = [result.summary_row() for result in bundle.predictions[:12]]
    if localization_rows:
        doc.add_heading("可能作用位置", level=2)
        _add_first_png(doc, dict(prediction_charts or {}), ("combined_",), "5.1", "Integrated protein localization prediction")
        _add_rows_table(doc, localization_rows, [("protein_id", "Protein"), ("tool", "Tool"), ("classification", "Classification"), ("provider", "Provider"), ("status", "Status")], [1800, 1300, 1800, 1800, 2660])
    else:
        _add_callout(doc, "定位状态", "本次未选择蛋白定位预测，或没有通过校验的蛋白序列。", fill=CAUTION_FILL)
    doc.add_heading("上游调控", level=2)
    _add_first_png(doc, visual_charts, ("promoter_regulation/",), "5.2", "Promoter TFBS distribution")
    candidate_rows = [{"rank": row.get("rank") or "—", "tf": row.get("tf") or "—", "tf_family": row.get("tf_family") or "—", "hit_count": row.get("hit_count") or "—", "best_p_value": f"{float(row.get('best_p_value')):.2g}" if row.get("best_p_value") not in (None, "") else "—", "nearest_tss_bp": row.get("nearest_tss_bp") or "—"} for row in bundle.upstream_tfs[:10]]
    _add_table_or_status(doc, candidate_rows, [("rank", "Rank"), ("tf", "TF/motif"), ("tf_family", "Family"), ("hit_count", "Hits"), ("best_p_value", "Best P"), ("nearest_tss_bp", "Nearest TSS")], [700, 2000, 1500, 900, 1500, 2760], "未选择启动子调控分析，或服务未返回候选 TF。")
    if bundle.mirna_targets:
        doc.add_heading("miRNA/RNAi Top 结果", level=2)
        _add_rows_table(doc, bundle.mirna_targets[:10], [("small_rna", "miRNA/sRNA"), ("target_transcript", "Target"), ("target_start", "Start"), ("expectation", "Expectation"), ("evidence_status", "Evidence boundary")], [1700, 2200, 1000, 1200, 3260])
    doc.add_heading("序列变异影响", level=2)
    _add_first_png(doc, visual_charts, ("variation/",), "5.3", "Variant distribution or haplotype summary")
    _add_table_or_status(doc, _variant_priority_rows(bundle.variants), [("position", "Position"), ("alleles", "REF → ALT"), ("region", "Region"), ("consequence", "Consequence"), ("af", "ALT AF"), ("qc", "REF/QC")], [1100, 1300, 1800, 2860, 1000, 1300], "未取得可解析变异；没有样本基因型矩阵时不推断单倍型。")
    haplotype_columns = [
        ("haplotype", "Haplotype"), ("sample_count", "Samples"),
        ("sample_frequency", "Frequency"), ("filtered_variant_count", "Variants"),
        ("subgroup_frequency", "Groups"),
    ]
    doc.add_heading("单倍型汇总", level=2)
    _add_table_or_status(doc, bundle.haplotypes[:20], haplotype_columns, [1800, 1100, 1200, 1100, 4160], "未形成单倍型汇总。")
    doc.add_heading("单倍型科研解读", level=2)
    _add_interpretation_blocks(doc, bundle, {"haplotype", "ai_haplotype"})
    _add_callout(doc, "注意事项", "TFBS、定位、miRNA/RNAi 和大多数变异影响均为计算预测；需结合独立实验、群体材料及 assembly/REF 一致性验证。", fill=CAUTION_FILL)

    _start_chapter(doc, "6", "综合科研判断与验证优先级", "把证据等级、冲突、关键缺口和实验优先级集中归纳，并回链到前述章节。")
    conflict_text = "；".join(serious_warnings[:4]) or "未记录会直接改变结论的 assembly/REF/映射冲突。"
    _add_key_value_table(doc, [
        ("已有证据", f"第2章：数据库/人工遗传证据 {len(bundle.genetic_evidence)} 条；关联文献 {len(bundle.literature_rows)} 篇。"),
        ("计算支持", f"第3-5章：表达记录 {len(bundle.efp_rows)}，结构域 {len(bundle.protein_domains)}，TFBS {len(bundle.promoter_tfbs)}，变异 {len(bundle.variants)}。"),
        ("证据冲突", conflict_text),
        ("关键缺口", "优先补齐未核验全文、缺失实验定位/互作、缺少功能互补或等位基因材料，以及无样本 GT 时的单倍型证据。"),
        ("推荐实验", "先验证与直接文献证据、表达场景和结构/变异线索同时一致的假设；随后开展遗传互补、亚细胞定位、酶活/互作及目标处理下表型验证。"),
    ])
    _add_callout(doc, "证据分级", "已有数据库/文献证据 ≠ 计算预测；计算支持 ≠ 因果结论；合理推测必须由明确实验验证。所有建议仅由本报告前述记录汇总，不新增无来源机制。", fill=CAUTION_FILL)
    _add_interpretation_blocks(doc, bundle, {"overall", "ai_overall", "ai_integrated"})

    _start_chapter(doc, "7", "方法、来源与警告", "集中记录数据库版本、检索日期、参数、失败服务和 assembly 边界，供复现与审阅。")
    _add_callout(doc, "关键结论", f"共记录 {len(bundle.sources)} 个数据/服务来源和 {len(bundle.warnings)} 条警告。一般警告集中在本章；正文仅保留会改变解释的边界提示。")
    option_rows = [{"item": key, "value": _compact_text(value, 150)} for key, value in bundle.analysis_options.items()]
    _add_table_or_status(doc, option_rows, [("item", "参数"), ("value", "设置")], [2700, 6660], "未记录额外运行参数。")
    doc.add_heading("警告", level=2)
    if bundle.warnings:
        for warning in bundle.warnings:
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.add_run(warning)
            _format_paragraph_runs(paragraph)
    else:
        _add_callout(doc, "状态", "未记录额外警告。")
    doc.add_heading("数据与服务来源", level=2)
    for source in bundle.sources or ["未记录来源"]:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(source)
        _format_paragraph_runs(paragraph)

    doc.add_page_break()
    doc.add_heading("附录", level=1)
    doc.add_heading("A1. eFP 数据源词典", level=2)
    glossary_rows = [{"source": key, "label": EFP_DATA_SOURCES.get(key, key), "display_label": efp_source_display_label(key), **EFP_SOURCE_GLOSSARY.get(key, {})} for key in EFP_DATA_SOURCES]
    _add_rows_table(doc, glossary_rows, [("display_label", "数据源"), ("scope", "组织/处理范围"), ("scale", "尺度"), ("id_namespace", "提交 ID")], [2000, 3800, 2460, 1100])
    doc.add_heading("A1.1 eFP 数据源适用问题与解读边界", level=3)
    _add_rows_table(
        doc,
        glossary_rows,
        [("display_label", "数据源"), ("design", "实验设计"), ("best_for", "适合回答"), ("caution", "解读边界")],
        [1500, 2500, 2580, 2780],
    )
    doc.add_heading("A1.2 eFP 官方来源与重复/汇总结构", level=3)
    _add_rows_table(
        doc,
        glossary_rows,
        [("display_label", "数据源"), ("reference", "官方来源/论文"), ("replicate_note", "重复或汇总结构")],
        [1800, 3460, 4100],
    )
    doc.add_heading("A2. 完整文献与遗传证据", level=2)
    reference_rows = [{"reference_id": row.get("reference_id"), "doi": row.get("doi"), "pmid": row.get("pmid"), "year": row.get("year"), "title": _compact_text(row.get("title"), 115), "verification_status": row.get("verification_status")} for row in bundle.ricedata_references]
    _add_table_or_status(doc, reference_rows, [("reference_id", "Ref ID"), ("doi", "DOI"), ("pmid", "PMID"), ("year", "Year"), ("title", "Title"), ("verification_status", "Verification")], [800, 1500, 1000, 700, 3560, 1800], "本次未解析到 RiceData 关联文献。")
    if include_full_sequences:
        doc.add_heading("A3. 完整序列", level=2)
        for sequence_type in SEQUENCE_TYPES:
            fasta = sequence_records_to_fasta(bundle.sequences, sequence_type)
            if not fasta:
                continue
            doc.add_heading(sequence_type, level=3)
            for line in fasta.rstrip().splitlines():
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                _format_run(paragraph.add_run(line), size=7.5)

    for paragraph in doc.paragraphs:
        if not paragraph.runs:
            paragraph.add_run("")
        if paragraph.style.name == "Heading 1":
            _format_paragraph_runs(paragraph, size=HEADING1_SIZE, bold=True)
        elif paragraph.style.name == "Heading 2":
            _format_paragraph_runs(paragraph, size=HEADING2_SIZE, bold=True)
        elif paragraph.style.name == "Title":
            _format_paragraph_runs(paragraph, size=TITLE_SIZE, bold=True)
        else:
            for run in paragraph.runs:
                if run.font.size is None:
                    _format_run(run)
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def _sheet_from_rows(workbook: Workbook, name: str, rows: list[dict[str, object]], columns: list[str]) -> None:
    sheet = workbook.create_sheet(name)
    sheet.append(columns)
    for row in rows:
        values = []
        for column in columns:
            value = row.get(column, "")
            if isinstance(value, (dict, list, tuple, set)):
                value = json.dumps(value, ensure_ascii=False, default=str)
            values.append(value)
        sheet.append(values)
    for cell in sheet[1]:
        cell.font = Font(name=EXCEL_FONT, bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="0F766E")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for row in sheet.iter_rows(min_row=2):
        for cell in row:
            cell.font = Font(name=EXCEL_FONT, size=10)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    for index, column in enumerate(columns, start=1):
        values = [str(sheet.cell(row=row, column=index).value or "") for row in range(1, min(sheet.max_row, 200) + 1)]
        sheet.column_dimensions[get_column_letter(index)].width = min(48, max(12, max(map(len, values), default=12) + 2))
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions


def build_excel_report(bundle: AnalysisBundle) -> bytes:
    workbook = Workbook()
    workbook.remove(workbook.active)
    overview = [
        {"item": "app_version", "value": "1.9.1 (build 20)"},
        {"item": "mode", "value": bundle.mode},
        {"item": "input_type", "value": bundle.input_type},
        {"item": "input_count", "value": len(bundle.inputs)},
        {"item": "sequence_count", "value": len(bundle.sequences)},
        {"item": "prediction_count", "value": len(bundle.predictions)},
        {"item": "ricedata_count", "value": len(bundle.ricedata_rows)},
        {"item": "efp_expression_count", "value": len(bundle.efp_rows)},
        {"item": "lab_omics_dataset_count", "value": len(bundle.lab_omics_datasets)},
        {"item": "lab_omics_differential_count", "value": len(bundle.lab_omics_differential)},
        {"item": "lab_omics_profile_count", "value": len(bundle.lab_omics_profiles)},
        {"item": "interpretation_requested_mode", "value": bundle.interpretation_status.get("requested_mode", "rules")},
        {"item": "interpretation_effective_mode", "value": bundle.interpretation_status.get("effective_mode", "rules")},
        {"item": "interpretation_provider", "value": bundle.interpretation_status.get("provider", "")},
        {"item": "interpretation_provider_label", "value": bundle.interpretation_status.get("provider_label", "")},
        {"item": "interpretation_model", "value": bundle.interpretation_status.get("model", "")},
        {"item": "interpretation_client_version", "value": bundle.interpretation_status.get("client_version", "")},
        {"item": "interpretation_error_code", "value": bundle.interpretation_status.get("error_code", "")},
        {"item": "interpretation_error", "value": bundle.interpretation_status.get("error", "")},
        {"item": "generated_at", "value": bundle.generated_at},
    ]
    _sheet_from_rows(workbook, "Overview", overview, ["item", "value"])
    interpretation_rows = _interpretation_rows(bundle)
    interpretation_columns = [
        "section", "title", "interpretation", "evidence_basis", "evidence_level",
        "confidence", "limitations", "recommended_action", "source_refs",
    ]
    _sheet_from_rows(workbook, "Interpretation", interpretation_rows, interpretation_columns)
    mapping_columns = ["input_id", "input_type", "resolved_rap_gene", "resolved_msu_id", "mapping_count", "status", "note", "error"]
    _sheet_from_rows(workbook, "ID_Mapping", bundle.mapping_rows, mapping_columns)
    ricedata_columns = list(bundle.ricedata_rows[0].keys()) if bundle.ricedata_rows else ["check", "status", "error"]
    _sheet_from_rows(workbook, "RiceData", bundle.ricedata_rows, ricedata_columns)
    reference_columns = list(bundle.ricedata_references[0].keys()) if bundle.ricedata_references else ["reference_id", "doi", "pmid", "title", "year", "verification_status", "status", "error"]
    _sheet_from_rows(workbook, "RiceData_References", bundle.ricedata_references, reference_columns)
    efp_rows = [record.summary_row() for record in bundle.efp_rows]
    efp_columns = list(efp_rows[0].keys()) if efp_rows else [
        "input_id", "msu_locus", "data_source", "tissue", "expression_level", "standard_deviation", "status", "error"
    ]
    _sheet_from_rows(workbook, "eFP_Expression", efp_rows, efp_columns)
    top_rows = expression_top_rows(bundle.efp_rows, limit=3)
    top_columns = list(top_rows[0].keys()) if top_rows else [
        "msu_locus", "data_source", "rank", "tissue", "expression_level", "standard_deviation"
    ]
    _sheet_from_rows(workbook, "eFP_Top_Tissues", top_rows, top_columns)
    glossary_rows = [{"data_source": key, "label": EFP_DATA_SOURCES.get(key, key), **EFP_SOURCE_GLOSSARY.get(key, {})} for key in EFP_DATA_SOURCES]
    _sheet_from_rows(
        workbook,
        "eFP_Source_Glossary",
        glossary_rows,
        [
            "data_source", "name_zh", "label", "scope", "design", "scale", "id_namespace",
            "reference", "replicate_note", "best_for", "outputs", "caution",
        ],
    )
    sequence_columns = list(next(iter([record.summary_row() for record in bundle.sequences]), {
        "input_id": "", "resolved_rap_gene": "", "resolved_msu_id": "", "transcript_id": "", "sequence_type": "", "length": "", "source": "", "assembly": "", "coordinates": "", "strand": "", "status": "", "validation_note": ""
    }).keys())
    _sheet_from_rows(workbook, "Sequence_Summary", [record.summary_row() for record in bundle.sequences], sequence_columns)
    sequence_plot_columns = list(bundle.sequence_plot_rows[0].keys()) if bundle.sequence_plot_rows else ["input_id", "rap_gene", "msu_id", "sequence_type", "length", "source", "assembly", "translation_consistency"]
    _sheet_from_rows(workbook, "Sequence_Plot_Data", bundle.sequence_plot_rows, sequence_plot_columns)
    prediction_columns = [
        "protein_id", "tool", "version", "status", "classification", "summary", "parameters",
        "region_count", "provider", "provider_job_id", "fallback_used", "result_url", "error",
    ]
    _sheet_from_rows(workbook, "Prediction_Summary", [result.summary_row() for result in bundle.predictions], prediction_columns)
    score_rows = [row for result in bundle.predictions for row in result.probability_rows()]
    _sheet_from_rows(
        workbook,
        "Prediction_Scores",
        score_rows,
        ["protein_id", "tool", "provider", "label", "probability"],
    )
    region_rows = [row for result in bundle.predictions for row in result.region_rows()]
    region_columns = ["protein_id", "tool", "region_type", "start", "end", "score", "sequence", "note"]
    _sheet_from_rows(workbook, "Prediction_Regions", region_rows, region_columns)
    deep_sheets = [
        ("Protein_Domains", bundle.protein_domains), ("Functional_Sites", bundle.functional_sites),
        ("Transcript_Models", bundle.transcript_models), ("Gene_Features", bundle.gene_features),
        ("Promoter_TFBS", bundle.promoter_tfbs), ("Upstream_TF", bundle.upstream_tfs),
        ("Variants", bundle.variants), ("Haplotype_Summary", bundle.haplotypes),
        ("miRNA_Targets", bundle.mirna_targets), ("RNAi_Offtargets", bundle.rnai_offtargets),
        ("Literature", bundle.literature_rows), ("Genetic_Evidence", bundle.genetic_evidence),
        ("Lab_Omics_Datasets", bundle.lab_omics_datasets),
        ("Lab_Omics_Comparisons", bundle.lab_omics_comparisons),
        ("Lab_Omics_Samples", bundle.lab_omics_samples),
        ("Lab_Omics_Differential", bundle.lab_omics_differential),
        ("Lab_Omics_Profiles", bundle.lab_omics_profiles),
        ("Lab_Omics_Status", bundle.lab_omics_status),
    ]
    for sheet_name, rows in deep_sheets:
        columns = list(rows[0].keys()) if rows else ["input_id", "status", "source_url", "queried_at", "error"]
        _sheet_from_rows(workbook, sheet_name, rows, columns)
    warning_rows = [{"type": "warning", "message": item} for item in bundle.warnings]
    warning_rows.extend(
        {"type": "prediction_consistency", "message": item}
        for item in prediction_consistency(bundle.predictions)
    )
    warning_rows.extend({"type": "source", "message": item} for item in bundle.sources)
    _sheet_from_rows(workbook, "Warnings_Sources", warning_rows, ["type", "message"])
    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def _rows_to_csv_bytes(rows: list[dict[str, object]]) -> bytes:
    if not rows:
        return b""
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=list(rows[0].keys()), extrasaction="ignore")
    writer.writeheader()
    writer.writerows(rows)
    return output.getvalue().encode("utf-8-sig")


def build_analysis_zip(
    bundle: AnalysisBundle,
    stem: str,
    word_bytes: bytes,
    excel_bytes: bytes,
    efp_charts: dict[str, bytes] | None = None,
    prediction_charts: dict[str, bytes] | None = None,
    prediction_raw_artifacts: dict[str, bytes] | None = None,
    deep_charts: dict[str, bytes] | None = None,
    deep_raw_artifacts: dict[str, bytes] | None = None,
) -> bytes:
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(f"{stem}.docx", word_bytes)
        archive.writestr(f"{stem}.xlsx", excel_bytes)
        archive.writestr("annotations/ricedata_gene_annotations.csv", _rows_to_csv_bytes(bundle.ricedata_rows))
        archive.writestr("annotations/ricedata_references.csv", _rows_to_csv_bytes(bundle.ricedata_references))
        efp_rows = [record.summary_row() for record in bundle.efp_rows]
        archive.writestr("expression/efp_expression_values.csv", _rows_to_csv_bytes(efp_rows))
        archive.writestr(
            "expression/source_metadata.json",
            json.dumps(
                {
                    "source": EFP_URL,
                    "guide": EFP_GUIDE_URL,
                    "mode": "Absolute",
                    "data_sources": list(dict.fromkeys(record.data_source for record in bundle.efp_rows)),
                    "data_source_glossary": EFP_SOURCE_GLOSSARY,
                    "note": "Official eFP values are preserved in the source-specific scale without cross-dataset rescaling. Absolute is not fold change.",
                },
                ensure_ascii=False,
                indent=2,
            ).encode("utf-8"),
        )
        for name, payload in sorted((efp_charts or {}).items()):
            archive.writestr(f"expression/figures/{name}", payload)
        for sequence_type in SEQUENCE_TYPES:
            fasta = sequence_records_to_fasta(bundle.sequences, sequence_type)
            filename = safe_file_stem(sequence_type.lower().replace("′", ""), "sequences") + ".fasta"
            archive.writestr(f"sequences/{filename}", fasta.encode("utf-8"))
        archive.writestr("sequences/sequence_relationship_plot_data.csv", _rows_to_csv_bytes(bundle.sequence_plot_rows))
        for index, result in enumerate(bundle.predictions, start=1):
            prefix = f"predictions/{safe_file_stem(result.protein_id)}/{safe_file_stem(result.tool)}_{index}"
            if result.raw_text:
                archive.writestr(prefix + ".txt", result.raw_text.encode("utf-8"))
            if result.raw_html:
                archive.writestr(prefix + ".html", result.raw_html.encode("utf-8"))
            if result.error:
                archive.writestr(prefix + "_error.txt", result.error.encode("utf-8"))
        for name, payload in sorted((prediction_charts or {}).items()):
            archive.writestr(f"predictions/figures/{safe_file_stem(name, 'figure')}", payload)
        for name, payload in sorted((prediction_raw_artifacts or {}).items()):
            parts = [safe_file_stem(part, "artifact") for part in Path(name).parts if part not in {"/", ".", ".."}]
            archive.writestr("predictions/raw/" + "/".join(parts), payload)
        trace_rows = [row for result in bundle.predictions for row in result.attempt_rows()]
        archive.writestr(
            "predictions/provider_trace.json",
            json.dumps(trace_rows, ensure_ascii=False, indent=2).encode("utf-8"),
        )
        deep_exports = {
            "protein_domains/protein_domains.csv": bundle.protein_domains,
            "protein_domains/functional_sites.csv": bundle.functional_sites,
            "gene_structure/transcript_models.csv": bundle.transcript_models,
            "gene_structure/gene_features.csv": bundle.gene_features,
            "promoter_regulation/promoter_tfbs.csv": bundle.promoter_tfbs,
            "promoter_regulation/upstream_tf.csv": bundle.upstream_tfs,
            "variation/variants.csv": bundle.variants,
            "variation/haplotype_summary.csv": bundle.haplotypes,
            "mirna_rnai/mirna_targets.csv": bundle.mirna_targets,
            "mirna_rnai/rnai_offtargets.csv": bundle.rnai_offtargets,
            "literature_evidence/literature.csv": bundle.literature_rows,
            "literature_evidence/genetic_evidence.csv": bundle.genetic_evidence,
            "literature_evidence/ricedata_references.csv": bundle.ricedata_references,
            "lab_omics/datasets.csv": bundle.lab_omics_datasets,
            "lab_omics/comparisons.csv": bundle.lab_omics_comparisons,
            "lab_omics/samples.csv": bundle.lab_omics_samples,
            "lab_omics/differential_results.csv": bundle.lab_omics_differential,
            "lab_omics/abundance_profiles.csv": bundle.lab_omics_profiles,
            "lab_omics/status.csv": bundle.lab_omics_status,
        }
        for name, rows in deep_exports.items():
            archive.writestr(name, _rows_to_csv_bytes(rows))
        for name, payload in sorted((deep_charts or {}).items()):
            archive.writestr(name.replace("/", "/figures/", 1), payload)
        for name, payload in sorted((deep_raw_artifacts or {}).items()):
            archive.writestr(name.replace("/", "/raw/", 1), payload)
        archive.writestr("deep_analysis/parameters.json", json.dumps(bundle.analysis_options, ensure_ascii=False, indent=2).encode("utf-8"))
        archive.writestr("deep_analysis/sources.json", json.dumps(bundle.sources, ensure_ascii=False, indent=2).encode("utf-8"))
        interpretation_rows = _interpretation_rows(bundle)
        archive.writestr("interpretation/interpretation.csv", _rows_to_csv_bytes(interpretation_rows))
        archive.writestr("interpretation/interpretation.json", json.dumps(interpretation_rows, ensure_ascii=False, indent=2).encode("utf-8"))
        archive.writestr("interpretation/status.json", json.dumps(bundle.interpretation_status, ensure_ascii=False, indent=2).encode("utf-8"))
        archive.writestr(
            "manifest.json",
            json.dumps(bundle.as_dict(), ensure_ascii=False, indent=2).encode("utf-8"),
        )
        archive.writestr(
            "README.txt",
            (
                "My Bio Tools v1.9.1 (build 20) - rice gene analysis bundle\n"
                f"Generated: {bundle.generated_at}\n"
                "Word fonts: East Asia=华文仿宋; ASCII/HAnsi=Times New Roman.\n"
                "All localization outputs are computational predictions and require experimental validation.\n"
            ).encode("utf-8"),
        )
    return output.getvalue()


def build_report_artifacts(
    bundle: AnalysisBundle,
    primary_name: str,
    efp_charts: dict[str, bytes] | None = None,
    prediction_charts: dict[str, bytes] | None = None,
    prediction_raw_artifacts: dict[str, bytes] | None = None,
    deep_charts: dict[str, bytes] | None = None,
    deep_raw_artifacts: dict[str, bytes] | None = None,
) -> dict[str, bytes | str]:
    if not bundle.generated_at:
        bundle.generated_at = datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds")
    if not bundle.interpretations:
        bundle.interpretations = build_rule_interpretations(bundle)
    if not bundle.interpretation_status:
        bundle.interpretation_status = {
            "requested_mode": "rules",
            "effective_mode": "rules",
            "rule_section_count": len(bundle.interpretations),
            "ai_section_count": 0,
            "error": "",
            "privacy": "离线规则解读未向外部模型发送数据。",
        }
    stem = (
        f"rice_gene_analysis_{safe_file_stem(primary_name)}"
        if bundle.mode == "单基因深度分析"
        else f"rice_gene_batch_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    )
    word_bytes = build_word_report(
        bundle,
        include_full_sequences=bundle.mode == "单基因深度分析",
        efp_charts=efp_charts,
        prediction_charts=prediction_charts,
        deep_charts=deep_charts,
    )
    excel_bytes = build_excel_report(bundle)
    zip_bytes = build_analysis_zip(
        bundle,
        stem,
        word_bytes,
        excel_bytes,
        efp_charts=efp_charts,
        prediction_charts=prediction_charts,
        prediction_raw_artifacts=prediction_raw_artifacts,
        deep_charts=deep_charts,
        deep_raw_artifacts=deep_raw_artifacts,
    )
    return {
        "stem": stem,
        "docx": word_bytes,
        "xlsx": excel_bytes,
        "zip": zip_bytes,
    }


__all__ = [
    "CHINESE_FONT",
    "EXCEL_FONT",
    "WESTERN_FONT",
    "build_analysis_zip",
    "build_excel_report",
    "build_report_artifacts",
    "build_word_report",
]
