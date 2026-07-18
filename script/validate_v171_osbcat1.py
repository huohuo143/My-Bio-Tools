#!/usr/bin/env python3
"""Live v1.7.1 acceptance case for Os03t0106400-01 / OsBCAT1."""

from __future__ import annotations

from datetime import datetime, timezone
import io
import json
from pathlib import Path
import sys
import zipfile

from docx import Document
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parent.parent
APP_SOURCE = ROOT / "app_source"
sys.path.insert(0, str(APP_SOURCE))

from analysis_jobs import RiceGeneAnalysisRequest  # noqa: E402
from rice_gene_analysis import INPUT_ID, MODE_SINGLE, execute_analysis_request  # noqa: E402
from rice_gene_core import SEQUENCE_TYPES  # noqa: E402
from rice_utr_promoter_downloader import TRANSCRIPT_SCOPE_CANONICAL  # noqa: E402


class Reporter:
    def update(self, phase, completed, total, detail=""):
        print(f"[{phase}] {detail}")

    def complete(self, phase, detail="", warning=False):
        print(f"[{phase}] {'WARNING ' if warning else ''}{detail}")

    def update_item(self, phase, item_key, completed, total, detail="", warning=False):
        print(f"[{phase}/{item_key}] {detail}")

    def complete_item(self, phase, item_key, detail="", warning=False):
        print(f"[{phase}/{item_key}] {'WARNING ' if warning else ''}{detail}")

    def is_cancelled(self):
        return False


def main() -> int:
    request = RiceGeneAnalysisRequest(
        project_name="v1.7.1 OsBCAT1 acceptance",
        mode=MODE_SINGLE,
        input_type=INPUT_ID,
        text="Os03t0106400-01",
        selected_types=tuple(SEQUENCE_TYPES),
        promoter_length=2000,
        transcript_scope=TRANSCRIPT_SCOPE_CANONICAL,
        selected_predictors=(),
        signalp_mode="fast",
        cnls_cutoff=5.0,
        nlstradamus_model=2,
        nlstradamus_cutoff=0.6,
        max_workers=3,
        include_ricedata=True,
        ricedata_depth="full",
        include_efp=True,
        efp_data_sources=("rice_rma", "ricestress_rma", "rice_single_cell"),
        selected_deep_analyses=("protein_domains", "gene_structure", "promoter_regulation", "literature_evidence"),
    )
    bundle, artifacts = execute_analysis_request(request, Reporter())
    output_dir = ROOT / "analysis_results" / "v1.7.1_osbcat1_acceptance"
    output_dir.mkdir(parents=True, exist_ok=True)
    stem = str(artifacts["stem"])
    docx_path = output_dir / f"{stem}.docx"
    xlsx_path = output_dir / f"{stem}.xlsx"
    zip_path = output_dir / f"{stem}.zip"
    docx_path.write_bytes(artifacts["docx"])
    xlsx_path.write_bytes(artifacts["xlsx"])
    zip_path.write_bytes(artifacts["zip"])

    references = {str(row.get("doi") or "").casefold(): row for row in bundle.ricedata_references}
    direct = references.get("10.1111/nph.16800", {})
    related = references.get("10.1111/nph.19551", {})
    evidence = next((row for row in bundle.genetic_evidence if row.get("evidence_type") == "mutation"), {})
    sequence_names = set(artifacts.get("deep_charts", {}))
    document = Document(io.BytesIO(artifacts["docx"]))
    heading_text = [paragraph.text for paragraph in document.paragraphs if paragraph.style.name.startswith("Heading")]
    expected_chapters = [
        "1. 基因概览与 ID 映射", "2. 已知功能、突变体与关联文献", "3. 表达模式与生物学场景",
        "4. 序列、转录本与蛋白结构", "5. 定位、调控与遗传变异", "6. 综合科研判断与验证优先级",
        "7. 方法、来源与警告", "附录",
    ]
    multi_column_tables = [table for table in document.tables if len(table.columns) > 1]
    workbook = load_workbook(io.BytesIO(artifacts["xlsx"]), read_only=True)
    with zipfile.ZipFile(io.BytesIO(artifacts["zip"])) as archive:
        zip_names = set(archive.namelist())

    single_cell_rows = [record for record in bundle.efp_rows if record.data_source == "rice_single_cell"]
    checks = {
        "identity_resolved": any(row.get("resolved_rap_gene") == "Os03g0106400" for row in bundle.mapping_rows),
        "symbol_is_osbcat1": any(row.get("GeneSymbol") == "OsBCAT1" for row in bundle.ricedata_rows),
        "mutant_evidence_present": "osbcat1" in str(evidence.get("evidence_text") or "").casefold(),
        "2020_direct_support": direct.get("verification_status") == "直接支持" and evidence.get("linked_dois") == "10.1111/nph.16800",
        "2020_pmid_completed": str(direct.get("pmid") or "") == "32654152",
        "2024_related_needs_verification": "需核验" in str(related.get("verification_status") or ""),
        "2024_pmid_completed": str(related.get("pmid") or "") == "38258425",
        "single_cell_submitted_rap": bool(single_cell_rows) and all(row.submitted_id == "Os03g0106400" and row.id_namespace == "RAP" for row in single_cell_rows),
        "sequence_svg_pdf_png": all(any(name.startswith("sequence_structure/") and name.endswith(extension) for name in sequence_names) for extension in (".svg", ".pdf", ".png")),
        "sequence_csv_present": "sequences/sequence_relationship_plot_data.csv" in zip_names,
        "all_new_chapters_present": all(chapter in heading_text for chapter in expected_chapters),
        "legacy_chapters_absent": not any(text.startswith("12. 文献") or text.startswith("14. 来源") for text in heading_text),
        "no_header_only_multicolumn_tables": all(len(table.rows) >= 2 for table in multi_column_tables),
        "excel_new_sheets": {"RiceData_References", "eFP_Source_Glossary", "Sequence_Plot_Data"}.issubset(workbook.sheetnames),
        "zip_complete": {"annotations/ricedata_references.csv", "literature_evidence/genetic_evidence.csv", "literature_evidence/ricedata_references.csv"}.issubset(zip_names),
    }
    report = {
        "validated_at_utc": datetime.now(timezone.utc).isoformat(),
        "app_version": "1.7.1",
        "build": 15,
        "input": "Os03t0106400-01",
        "gene_symbol": "OsBCAT1",
        "success": all(checks.values()),
        "checks": checks,
        "counts": {
            "sequences": len(bundle.sequences), "efp_rows": len(bundle.efp_rows), "literature": len(bundle.literature_rows),
            "genetic_evidence": len(bundle.genetic_evidence), "ricedata_references": len(bundle.ricedata_references),
            "warnings": len(bundle.warnings),
        },
        "references": bundle.ricedata_references,
        "warnings": bundle.warnings,
        "artifacts": {"docx": str(docx_path), "xlsx": str(xlsx_path), "zip": str(zip_path)},
    }
    report_path = output_dir / "v1.7.1_osbcat1_acceptance_report.json"
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2, default=str) + "\n", encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2, default=str))
    return 0 if report["success"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
