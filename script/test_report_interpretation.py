#!/usr/bin/env python3
"""Focused regression tests for report interpretation and export delivery."""

from __future__ import annotations

import io
import json
from pathlib import Path
import sys
import unittest
import zipfile

from docx import Document
from openpyxl import load_workbook

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "app_source"))

from report_builder import build_report_artifacts  # noqa: E402
from report_interpretation import (  # noqa: E402
    MODE_LLM,
    PROVIDER_OLLAMA,
    PROVIDER_OPENAI_COMPATIBLE,
    build_rule_interpretations,
    generate_interpretations,
    probe_model_connection,
)
from rice_gene_core import AnalysisBundle  # noqa: E402


class _Response:
    def __init__(self, payload: dict[str, object]):
        self.payload = payload

    def raise_for_status(self) -> None:
        return None

    def json(self) -> dict[str, object]:
        return self.payload


class _Session:
    def __init__(self, response: dict[str, object] | None = None, error: Exception | None = None):
        self.response = response or {}
        self.error = error
        self.calls: list[tuple[str, dict[str, object]]] = []

    def post(self, url: str, **kwargs: object) -> _Response:
        self.calls.append((url, kwargs))
        if self.error:
            raise self.error
        return _Response(self.response)


class ReportInterpretationTests(unittest.TestCase):
    def test_model_connection_probes_ollama_and_compatible_api(self) -> None:
        ollama = _Session({"message": {"content": "OK"}})
        detail = probe_model_connection(
            provider=PROVIDER_OLLAMA,
            base_url="http://127.0.0.1:11434",
            model="qwen-fixture",
            session=ollama,
        )
        self.assertIn("qwen-fixture", detail)
        self.assertTrue(ollama.calls[0][0].endswith("/api/chat"))

        cloud = _Session({"choices": [{"message": {"content": "OK"}}]})
        detail = probe_model_connection(
            provider=PROVIDER_OPENAI_COMPATIBLE,
            base_url="https://api.example.test/v1",
            model="cloud-fixture",
            api_key="secret-test-key",
            session=cloud,
        )
        self.assertIn("cloud-fixture", detail)
        self.assertTrue(cloud.calls[0][0].endswith("/chat/completions"))
        self.assertEqual(cloud.calls[0][1]["headers"]["Authorization"], "Bearer secret-test-key")

    def fixture(self) -> AnalysisBundle:
        bundle = AnalysisBundle(
            mode="单基因深度分析",
            input_type="RAP/MSU ID",
            inputs=["LOC_Os01g01010"],
            generated_at="2026-07-18T18:00:00+08:00",
        )
        bundle.lab_omics_differential = [
            {
                "msu_locus": "LOC_Os01g01010", "dataset_name": "RSV RNA-seq",
                "comparison_name": "RSV vs control", "assay": "mRNA", "log2fc": 2.2,
                "padj": 0.01, "descriptive": False, "source_file": "/secret/project/raw.xlsx",
            },
            {
                "msu_locus": "LOC_Os01g01010", "dataset_name": "RSV proteome",
                "comparison_name": "RSV vs control", "assay": "protein", "log2fc": -0.8,
                "pvalue": 0.04, "descriptive": False,
            },
            {
                "msu_locus": "LOC_Os01g01010", "dataset_name": "RSV ubiquitome",
                "comparison_name": "RSV vs control", "assay": "ubiquitination", "log2fc": 3.1,
                "site_position": 123, "site_residue": "K", "descriptive": True,
            },
        ]
        bundle.variants = [{"position": 100}, {"position": 120}]
        bundle.haplotypes = [
            {
                "input_id": "LOC_Os01g01010", "haplotype": "H1", "sample_count": 8,
                "sample_frequency": 0.8, "subgroup_frequency": "Indica:6/7;Japonica:2/3",
                "filtered_variant_count": 2, "status": "calculated",
            },
            {
                "input_id": "LOC_Os01g01010", "haplotype": "H2", "sample_count": 2,
                "sample_frequency": 0.2, "subgroup_frequency": "Indica:1/7;Japonica:1/3",
                "filtered_variant_count": 2, "status": "calculated",
            },
        ]
        return bundle

    def test_rules_explain_omics_and_haplotype_boundaries(self) -> None:
        rows = build_rule_interpretations(self.fixture())
        omics = next(row for row in rows if row["section"] == "lab_omics")
        haplotype = next(row for row in rows if row["section"] == "haplotype")
        self.assertIn("log2FC=3.100", omics["interpretation"])
        self.assertIn("PTM", omics["limitations"])
        self.assertIn("80.0%", haplotype["interpretation"])
        self.assertIn("不等同于性状关联", haplotype["interpretation"])
        self.assertIn("1−Σp²", haplotype["evidence_basis"])

    def test_llm_mode_is_labeled_and_payload_is_redacted(self) -> None:
        content = json.dumps({
            "executive_summary": "结构化证据支持优先验证。",
            "multiomics_interpretation": "mRNA与蛋白方向不一致，需区分转录后调控。",
            "haplotype_interpretation": "主单倍型占优，但尚无性状关联。",
            "integrated_hypotheses": [],
            "gene_identity": {},
            "core_function": {},
            "mechanism_chains": [],
            "context_branches": [],
            "omics_integration": [],
            "testable_hypotheses": [],
            "knowledge_gaps": [],
            "references": [],
        }, ensure_ascii=False)
        session = _Session({"message": {"content": content}})
        rows, status = generate_interpretations(
            self.fixture(), mode=MODE_LLM, provider=PROVIDER_OLLAMA,
            base_url="http://127.0.0.1:11434", model="fixture", session=session,
        )
        self.assertEqual(status["effective_mode"], MODE_LLM)
        self.assertTrue(any(row["section"] == "ai_lab_omics" for row in rows))
        self.assertTrue(all(row["evidence_level"] == "AI辅助推断" for row in rows if str(row["section"]).startswith("ai_")))
        sent = json.dumps(session.calls, ensure_ascii=False)
        self.assertNotIn("/secret/project/raw.xlsx", sent)
        self.assertTrue(session.calls[0][0].endswith("/api/chat"))

    def test_llm_failure_falls_back_to_rules(self) -> None:
        rows, status = generate_interpretations(
            self.fixture(), mode=MODE_LLM, provider=PROVIDER_OLLAMA,
            base_url="http://127.0.0.1:11434", model="fixture",
            session=_Session(error=ConnectionError("offline")),
        )
        self.assertEqual(status["effective_mode"], "rules")
        self.assertIn("offline", status["error"])
        self.assertFalse(any(str(row["section"]).startswith("ai_") for row in rows))

    def test_word_excel_and_zip_contain_interpretation(self) -> None:
        bundle = self.fixture()
        bundle.interpretations = build_rule_interpretations(bundle)
        bundle.interpretation_status = {"requested_mode": "rules", "effective_mode": "rules", "error": ""}
        artifacts = build_report_artifacts(bundle, "LOC_Os01g01010")

        document = Document(io.BytesIO(artifacts["docx"]))
        text = "\n".join(
            [paragraph.text for paragraph in document.paragraphs]
            + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
        )
        self.assertIn("多组学科研解读", text)
        self.assertIn("单倍型科研解读", text)
        self.assertIn("不等同于性状关联", text)

        workbook = load_workbook(io.BytesIO(artifacts["xlsx"]), read_only=True)
        self.assertIn("Interpretation", workbook.sheetnames)
        self.assertGreater(workbook["Interpretation"].max_row, 1)

        with zipfile.ZipFile(io.BytesIO(artifacts["zip"])) as archive:
            names = set(archive.namelist())
            self.assertIn("interpretation/interpretation.csv", names)
            self.assertIn("interpretation/interpretation.json", names)
            self.assertIn("interpretation/status.json", names)


if __name__ == "__main__":
    unittest.main()
